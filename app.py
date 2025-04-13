import os
from flask import Flask, render_template, request, redirect, url_for, make_response
from datetime import datetime
from docx import Document
import funciones as fc

app = Flask(__name__)

# Ruta principal: mostrar formulario
@app.route('/')
def index():
    return render_template('index.html', resultado=None)

# Ruta para procesar el formulario
@app.route('/generar_demanda', methods=['GET','POST'])
def generar_demanda():
    if request.method == 'POST':

        try:
            # Obtener datos del formulario
            nombre_madre = request.form.get('nombre_madre')
            municipio_madre = request.form.get('municipio_madre')
            cedula_madre = request.form.get('cedula_madre')
            nombre_padre = request.form.get('nombre_padre')
            municipio_padre = request.form.get('municipio_padre')
            cedula_padre = request.form.get('cedula_padre')
            ciudad_menor = request.form.get('ciudad_menor')
            dia_audiencia = request.form.get('dia_audiencia')
            acuerdo = request.form.get('acuerdo')
            fecha_inicio = request.form.get('fecha_inicio')
            fecha_fin = request.form.get('fecha_fin')
            cuota_inicial = float(request.form.get('cuota_inicial', 0))

            # Validar fechas
            try:
                fecha_inicio_dt = datetime.strptime(fecha_inicio, '%Y-%m-%d')
                fecha_fin_dt = datetime.strptime(fecha_fin, '%Y-%m-%d')
            except ValueError:
                return render_template('index.html', error="Formato de fecha inválido. Use el selector de fechas.")

            if fecha_inicio_dt >= fecha_fin_dt:
                return render_template('index.html', error="La fecha de inicio debe ser anterior a la fecha de fin.")

            # Procesar niños (puede haber múltiples)
            niños = {}
            i = 1
            while f'nombre_nino_{i}' in request.form:
                niños[f'niño{i}'] = {
                    'nombre': request.form.get(f'nombre_nino_{i}'),
                    'identificacion': request.form.get(f'identificacion_nino_{i}')
                }
                i += 1

            # Generar contenido usando las funciones
            introduccion, nombres = fc.escribir_introduccion(
                nombre_madre, municipio_madre, cedula_madre,
                nombre_padre, municipio_padre, cedula_padre,
                ciudad_menor, niños
            )
            hecho_1 = fc.escribir_hecho_1(nombre_madre, nombre_padre, nombres, ciudad_menor)
            hecho_2 = fc.escribir_hecho_2(dia_audiencia, acuerdo)
            hecho_3 = fc.escribir_hecho_3(fecha_inicio, fecha_fin)
            cantidad_meses = fc.calcular_meses_entre_fechas(fecha_inicio, fecha_fin)
            df_cuotas = fc.crear_dataframe_cuotas_variables(fecha_inicio, cantidad_meses, cuota_inicial)
            hecho_4_texto, hecho_4_tabla = fc.escribir_hecho_4(nombre_padre, df_cuotas)

            # Crear el documento Word
            doc = Document()
            doc.add_heading('Demanda Ejecutiva de Alimentos', 0)
            doc.add_paragraph(introduccion)
            doc.add_paragraph(hecho_1)
            doc.add_paragraph(hecho_2)
            doc.add_paragraph(hecho_3)
            doc.add_paragraph(hecho_4_texto)
            table = doc.add_table(rows=1, cols=len(df_cuotas.columns))
            hdr_cells = table.rows[0].cells
            
            # Agregar encabezados a la tabla
            for i, column in enumerate(df_cuotas.columns):
                hdr_cells[i].text = column
            
            # Agregar filas a la tabla
            for _, row in df_cuotas.iterrows():
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)
            
            os.makedirs('static', exist_ok=True)
            doc.save('static/Demanda_Ejecutiva.docx')

        
            # Pasar las partes por separado al template
            return render_template('resultado.html', 
                                introduccion=introduccion,
                                hecho_1=hecho_1,
                                hecho_2=hecho_2,
                                hecho_3=hecho_3,
                                hecho_4_texto=hecho_4_texto,
                                hecho_4_tabla=hecho_4_tabla)
                                #hecho_4_tabla=df_cuotas.to_html(index=False),
                                #descargar=True)
        except Exception as e:
            return render_template('index.html', error=f"Error al procesar los datos: {str(e)}")
        
    else :
        return render_template('index.html', resultado=None)
        
if __name__ == '__main__':
    app.run(debug=True)